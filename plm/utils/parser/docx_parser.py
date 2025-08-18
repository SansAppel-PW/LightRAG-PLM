import asyncio
import os
import re
import subprocess
from collections import Counter

import pandas as pd
from docx import Document
from io import BytesIO
from plm.utils.tokenize import doc_tokenizer
from plm.utils.parser.base_parser import BaseParser
from plm.utils.exception_utils import exception_handler
from plm.utils.handler.exception_handler import DocxParserException
from loguru import logger


class ParserException(Exception):
    def __init__(self, doc_name, doc_path):
        self.doc_name = doc_name
        self.doc_path = doc_path

    def __str__(self):
        return f"doc_name:{self.doc_name}, doc_path:{self.doc_path}"


def do_parse(output_dir,
             docx_file_paths: list[str],
             docx_bytes_list: list[bytes],
             docx_file_names: list[str],
             d_lang_list: list[str],
             backend="pipeline",
             parse_method="auto",
             image_enable=True,
             table_enable=True,
             server_url=None,
             start_page_id=0,
             end_page_id=None,
             **kwargs,
             ):
    if backend == "pipeline":
        _process_pipeline(
            output_dir, docx_file_paths, docx_bytes_list, docx_file_names, d_lang_list,
            parse_method, image_enable, table_enable
        )
    else:
        # MinerU
        # Recognize Tables and Transform to HTML
        logger.error(f'PARSER DOCX BACKEND IS NOT PIPELINE...')
        # if backend.startswith("vlm-"):
        #     backend = backend[4:]
        #
        # os.environ['MINERU_VLM_FORMULA_ENABLE'] = str(formula_enable)
        # os.environ['MINERU_VLM_TABLE_ENABLE'] = str(table_enable)
        #
        # _process_vlm(
        #     output_dir, pdf_file_names, pdf_bytes_list, backend,
        #     f_draw_layout_bbox, f_draw_span_bbox, f_dump_md, f_dump_middle_json,
        #     f_dump_model_output, f_dump_orig_pdf, f_dump_content_list, f_make_md_mode,
        #     server_url, **kwargs,
        # )


def _process_pipeline(
        output_dir,
        docx_file_paths,
        docx_bytes_list,
        docx_file_names,
        d_lang_list,
        parse_method,
        image_enable,
        table_enable,
        # f_draw_layout_bbox,
        # f_draw_span_bbox,
        # f_dump_md,
        # f_dump_middle_json,
        # f_dump_model_output,
        # f_dump_orig_pdf,
        # f_dump_content_list,
        # f_make_md_mode,
):
    # """处理pipeline后端逻辑"""
    # from mineru.backend.pipeline.model_json_to_middle_json import \
    #     result_to_middle_json as pipeline_result_to_middle_json
    # from mineru.backend.pipeline.pipeline_analyze import doc_analyze as pipeline_doc_analyze
    #
    # infer_results, all_image_lists, all_pdf_docs, lang_list, ocr_enabled_list = (
    #     pipeline_doc_analyze(
    #         pdf_bytes_list, p_lang_list, parse_method=parse_method,
    #         formula_enable=p_formula_enable, table_enable=p_table_enable
    #     )
    # )
    if not docx_bytes_list:
        docx_bytes_list =[]
        if not docx_file_paths:
            for docx_file_path in docx_file_paths:
                with open(docx_file_path, 'rb') as f:
                    docx_bytes_list.append(f.read())

    for idx, doc_bytes in enumerate(docx_bytes_list):

        docx = Document(doc_bytes)
        model_json = copy.deepcopy(model_list)
        pdf_file_name = pdf_file_names[idx]
        local_image_dir, local_md_dir = prepare_env(output_dir, pdf_file_name, parse_method)
        image_writer, md_writer = FileBasedDataWriter(local_image_dir), FileBasedDataWriter(local_md_dir)

        images_list = all_image_lists[idx]
        pdf_doc = all_pdf_docs[idx]
        _lang = lang_list[idx]
        _ocr_enable = ocr_enabled_list[idx]

        middle_json = pipeline_result_to_middle_json(
            model_list, images_list, pdf_doc, image_writer,
            _lang, _ocr_enable, p_formula_enable
        )

        pdf_info = middle_json["pdf_info"]
        pdf_bytes = pdf_bytes_list[idx]

        _process_output(
            pdf_info, pdf_bytes, pdf_file_name, local_md_dir, local_image_dir,
            md_writer, f_draw_layout_bbox, f_draw_span_bbox, f_dump_orig_pdf,
            f_dump_md, f_dump_content_list, f_dump_middle_json, f_dump_model_output,
            f_make_md_mode, middle_json, model_json, is_pipeline=True
        )


class DocxParser(BaseParser):

    def __init__(self, *args, **kwargs):  # file_path: str,
        # path = kwargs['path']
        # self.file_path = file_path
        # self.doc = Document(file_path)
        # self.doc_name = os.path.splitext(os.path.basename(file_path))[0]
        # self.minio_bucket = 'plm'
        pass

    def __extract_table_content(self, tb):
        """

        :param tb:
        :return:
        """
        df = []
        # transform the type of data from Document.Table to Pandas.DataFrame
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        """

        :param df:
        :return:
        """

        def blockType(b):
            patt = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in patt:
                if re.search(p, b):
                    return n
            tks = [t for t in doc_tokenizer.tokenize(b).split() if len(t) > 1]
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"
            if len(tks) == 1 and doc_tokenizer.tag(tks[0]) == "nr":
                return "Nr"
            return "Ot"

        # the table should have at least two rows of data to distinguish between headers and data
        if len(df) < 2:
            return []

        # count the types of all cells starting from second row (Skipping the headers)
        max_type = Counter(
            [blockType(str(df.iloc[i, j])) for i in range(1, len(df)) for j in range(len(df.iloc[i, :]))])
        # find the most common type (e.g. Nu represents a numerical table)
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]  # header is not necessarily appear in the first line
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j])) for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)
        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            hr = [r - i for r in hdrows]  # relative positions of header rows
            hr = [r for r in hr if r < 0]  # only previous headers
            t = len(hr) - 1
            # find the consecutive nearest header rows by traversing in reverse order
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            headers = []
            for j in range(len(df.iloc[i, :])):  # row i in the table
                t = []
                for h in hr:  # multiple header row
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):  # row i in the table
                if not str(df.iloc[i, j]):  # column j in the table
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))
        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    # @staticmethod
    @exception_handler(
        retries=0,  # don't retry
        delay=1,
        exceptions=DocxParserException,
        default=None,
        log_traceback=True
    )
    def __call__(self, fnm, from_page=0, to_page=100000000):
        self.doc = Document(fnm) if isinstance(fnm, str) else Document(BytesIO(fnm))
        pn = 0  # parsed page
        secs = []  # parsed contents
        for p in self.doc.paragraphs:
            if pn > to_page:  # page num limitation
                break

            runs_within_single_paragraph = []  # save runs within the range of pages
            for run in p.runs:
                if pn > to_page:
                    break
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text)  # append run.text first
                if 'lastRenderedPageBreak' in run._element.xml:  # wrap page break checker into a static method
                    pn += 1
            # then concat run.text as part of the paragraph
            secs.append(("".join(runs_within_single_paragraph), p.style.name if hasattr(p.style, 'name') else ''))
        # 解析表格报错：nltk_tokenize find 文件路径在windows 下有问题，暂时只返回secs
        # tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        # tbls = self.doc.tables
        # return secs, tbls
        return ','.join(['<|SECTION|>'.join(map(str, sec)) for sec in secs])


if __name__ == "__main__":
    file_path = r"C:\Users\houzhimingwx1\Documents\01-code\00-hik-yf\Intelligent_QA\PLM2.0\assets\plm_docx\BOM 审核申请.docx"
    secs = DocxParser()(fnm=file_path)
    logger.info(f'Docx Parser Result: {secs}')
