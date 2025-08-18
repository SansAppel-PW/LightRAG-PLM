import os
import re
import json
import uuid
import logging
import tempfile
import shutil
import traceback
from collections import defaultdict, deque
from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image
import io
from lxml import etree
import sys
import zipfile

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("docx_parser.log"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


def iter_block_items(parent):
    """
    按文档顺序生成段落和表格
    """
    if isinstance(parent, _Document):
        parent_elem = parent.element.body
    elif isinstance(parent, Table):
        parent_elem = parent._tbl
    else:
        raise ValueError("Unsupported parent type")

    for child in parent_elem.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def clean_text(text):
    """清理文本中的特殊字符和多余空格"""
    if not text:
        return ""
    # 替换各种空格和特殊字符
    text = text.replace('\xa0', ' ').replace('\t', ' ').replace('\r', '')
    # 合并多余空格
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def get_heading_level(para):
    """
    智能识别标题层级
    1. 通过样式名称识别（如"Heading 1"或"标题1"）
    2. 通过文本模式识别（如"1.1 标题内容"）
    """
    # 通过样式名称识别
    if hasattr(para, 'style') and para.style and para.style.name:
        style_name = para.style.name.lower()
        if 'heading' in style_name or '标题' in style_name:
            # 提取样式中的数字
            match = re.search(r'\d+', style_name)
            if match:
                return int(match.group())
            return 1  # 默认一级标题

    # 通过文本模式识别 (如 "1.1 标题内容")
    text = clean_text(para.text)
    match = re.match(r'^(\d+(\.\d+)*)\s+', text)
    if match:
        level_str = match.group(1)
        return min(len(level_str.split('.')), 6)  # 最大支持6级

    return 0


def is_list_item(para):
    """判断段落是否为列表项"""
    if not hasattr(para, '_p') or para._p is None:
        return False
    if not hasattr(para._p, 'pPr') or para._p.pPr is None:
        return False
    return para._p.pPr.numPr is not None


def get_list_info(para, list_counter):
    """
    获取列表项详细信息
    返回: (list_level, prefix)
    """
    num_pr = para._p.pPr.numPr
    if num_pr is None:
        return 0, ""

    # 获取列表层级
    ilvl = num_pr.ilvl
    list_level = int(ilvl.val) if ilvl is not None and ilvl.val is not None else 0

    # 获取列表类型
    num_id = num_pr.numId
    num_id_val = num_id.val if num_id is not None else None

    # 更新列表计数器
    list_counter[list_level] += 1
    # 重置子层级计数器
    for lvl in range(list_level + 1, 10):
        if lvl in list_counter:
            list_counter[lvl] = 0

    # 创建前缀
    prefix = " " * (list_level * 4)  # 每级缩进4个空格

    # 确定列表符号
    is_bullet = "bullet" in str(para.style.name).lower()
    if is_bullet:
        prefix += "• "  # 项目符号
    else:
        # 数字编号
        prefix += ".".join(str(list_counter[lvl]) for lvl in range(list_level + 1)) + ". "

    return list_level, prefix


def extract_images_from_xml(xml_str, doc_part, images_dir, image_references, context=""):
    """
    从XML字符串中提取图片并保存
    返回: 图片节点列表
    """
    image_nodes = []
    try:
        if not xml_str or ('<pic:pic' not in xml_str and '<w:drawing>' not in xml_str):
            return image_nodes

        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }

        root = etree.fromstring(xml_str)

        # 查找所有图片元素
        for blip in root.findall('.//a:blip', namespaces):
            embed_id = blip.get(f'{{{namespaces["r"]}}}embed')
            if not embed_id:
                continue

            # 获取图片部件
            if embed_id not in doc_part.related_parts:
                logger.warning(f"图片关系 {embed_id} 在 {context} 中未找到")
                continue

            image_part = doc_part.related_parts[embed_id]
            image_data = image_part.blob

            # 确定图片格式
            img_format = "png"  # 默认格式
            if hasattr(image_part, 'content_type'):
                if 'jpeg' in image_part.content_type:
                    img_format = "jpg"
                elif 'gif' in image_part.content_type:
                    img_format = "gif"
                elif 'bmp' in image_part.content_type:
                    img_format = "bmp"
                elif 'png' in image_part.content_type:
                    img_format = "png"
                elif 'svg' in image_part.content_type:
                    img_format = "svg"
                elif 'tiff' in image_part.content_type:
                    img_format = "tiff"

            # 生成唯一图片ID
            image_id = f"img_{uuid.uuid4().hex[:8]}"
            img_filename = f"{image_id}.{img_format}"
            image_path = os.path.join(images_dir, img_filename)

            # 保存图片
            with open(image_path, "wb") as img_file:
                img_file.write(image_data)

            # 获取图片尺寸
            width, height = 0, 0
            try:
                with Image.open(io.BytesIO(image_data)) as img:
                    width, height = img.size
            except Exception as e:
                # 对于SVG等无法直接获取尺寸的图片格式，跳过尺寸获取
                pass

            # 创建图片节点
            image_node = {
                "type": "image",
                "url": f"images/{img_filename}",
                "format": img_format,
                "width": width,
                "height": height,
                "size": f"{len(image_data) / 1024:.2f} KB",
                "context": context
            }

            # 记录图片信息
            image_references[image_id] = image_node
            image_nodes.append(image_node)

    except Exception as e:
        logger.error(f"从XML提取图片失败: {e}")

    return image_nodes


def extract_paragraph_images(para, images_dir, image_references):
    """提取段落中的图片并返回图片节点列表"""
    image_nodes = []
    context = f"段落: {clean_text(para.text[:20])}..." if para.text else "段落"

    for run_idx, run in enumerate(para.runs):
        if run._element and run._element.xml:
            try:
                # 提取该运行中的图片
                images = extract_images_from_xml(
                    run._element.xml,
                    para.part,
                    images_dir,
                    image_references,
                    f"{context} (运行 {run_idx})"
                )
                if images:
                    image_nodes.extend(images)
            except Exception as e:
                logger.error(f"提取段落图片失败: {e}")

    return image_nodes


def extract_table_images(cell, table_idx, row_idx, cell_idx, image_references, images_dir):
    """提取表格单元格中的图片"""
    image_nodes = []
    context = f"表格{table_idx}单元格[{row_idx},{cell_idx}]"

    try:
        if not hasattr(cell, '_tc') or cell._tc is None:
            return image_nodes

        cell_xml = etree.tostring(cell._tc, encoding='unicode')
        image_nodes = extract_images_from_xml(
            cell_xml,
            cell.part,
            images_dir,
            image_references,
            context
        )
    except Exception as e:
        logger.error(f"提取表格图片失败: {e}")

    return image_nodes


def identify_merged_cells(table):
    """
    识别合并的单元格
    返回: 被合并单元格的位置集合{(row_idx, col_idx)}
    """
    merged_cells = set()
    try:
        grid = table._tbl.find('.//w:tblGrid', table._tbl.nsmap)
        if grid is None:
            return merged_cells

        cols = grid.findall('.//w:gridCol', grid.nsmap)
        if not cols:
            return merged_cells

        # 创建单元格网格
        rows = table.rows
        if not rows:
            return merged_cells

        # 标记所有被合并的单元格
        for row_idx, row in enumerate(rows):
            for cell_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.tcPr

                # 检测跨列
                grid_span = 1
                if tcPr.gridSpan is not None:
                    grid_span = int(tcPr.gridSpan.val)

                # 检测跨行
                row_span = 1
                v_merge = tcPr.vMerge
                if v_merge is not None and v_merge.val == "restart":
                    # 计算实际行跨度
                    for i in range(row_idx + 1, len(table.rows)):
                        if cell_idx >= len(table.rows[i].cells):
                            break
                        next_cell = table.rows[i].cells[cell_idx]
                        next_tcPr = next_cell._tc.tcPr
                        if next_tcPr.vMerge is None or next_tcPr.vMerge.val != "continue":
                            break
                        row_span += 1

                # 标记被合并的单元格
                if grid_span > 1 or row_span > 1:
                    for r in range(row_idx, row_idx + row_span):
                        for c in range(cell_idx, cell_idx + grid_span):
                            if r != row_idx or c != cell_idx:  # 跳过起始单元格
                                if r < len(rows) and c < len(rows[r].cells):
                                    merged_cells.add((r, c))
    except Exception as e:
        logger.error(f"识别合并单元格失败: {e}")

    return merged_cells


def parse_table(table, table_idx, image_references, images_dir):
    """解析表格并处理合并单元格"""
    # 识别所有合并单元格
    merged_cells = identify_merged_cells(table)

    table_data = []
    for row_idx, row in enumerate(table.rows):
        row_nodes = []
        for cell_idx, cell in enumerate(row.cells):
            # 如果是被合并的单元格，跳过
            if (row_idx, cell_idx) in merged_cells:
                continue

            # 创建单元格节点
            cell_node = {
                "type": "table_cell",
                "row": row_idx,
                "col": cell_idx,
                "content": []
            }

            # 添加文本内容
            cell_text = clean_text(cell.text)
            if cell_text:
                cell_node["content"].append({
                    "type": "text",
                    "text": cell_text
                })

            # 提取单元格中的图片
            image_nodes = extract_table_images(cell, table_idx, row_idx, cell_idx, image_references, images_dir)
            if image_nodes:
                for img in image_nodes:
                    cell_node["content"].append({
                        "type": "image",
                        "url": img["url"],
                        "format": img["format"],
                        "width": img["width"],
                        "height": img["height"],
                        "size": img["size"]
                    })

            row_nodes.append(cell_node)
        table_data.append(row_nodes)

    return table_data


def extract_metadata(doc, docx_path):
    """提取文档元数据"""
    core_props = doc.core_properties
    return {
        "source_path": docx_path,
        "title": core_props.title,
        "author": core_props.author,
        "created": str(core_props.created),
        "modified": str(core_props.modified),
        "subject": core_props.subject,
        "keywords": core_props.keywords,
        "category": core_props.category,
        "comments": core_props.comments,
        "company": getattr(core_props, "company", "N/A"),
        "file_size": f"{os.path.getsize(docx_path) / 1024:.2f} KB"
    }


def extract_header_footer_images(doc, images_dir, image_references):
    """提取页眉页脚中的图片"""
    try:
        for section in doc.sections:
            # 页眉
            if section.header:
                for paragraph in section.header.paragraphs:
                    images = extract_paragraph_images(paragraph, images_dir, image_references)
                    if images:
                        # 为页眉图片添加独立节点
                        section.header_part.images = section.header_part.images or []
                        section.header_part.images.extend(images)

            # 页脚
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    images = extract_paragraph_images(paragraph, images_dir, image_references)
                    if images:
                        # 为页脚图片添加独立节点
                        section.footer_part.images = section.footer_part.images or []
                        section.footer_part.images.extend(images)
    except Exception as e:
        logger.error(f"提取页眉页脚图片失败: {e}")


def parse_docx(docx_path, output_dir):
    """
    解析单个DOCX文档并提取内容
    返回结构化JSON数据
    """
    try:
        # 创建临时工作目录
        temp_dir = tempfile.mkdtemp(prefix="docx_extract_")
        logger.info(f"创建临时目录: {temp_dir}")

        # 复制文件到临时目录
        temp_docx_path = os.path.join(temp_dir, os.path.basename(docx_path))
        shutil.copy2(docx_path, temp_docx_path)

        # 检查文件是否成功复制
        if not os.path.exists(temp_docx_path) or os.path.getsize(temp_docx_path) == 0:
            logger.error(f"复制后的文件不存在或为空: {temp_docx_path}")
            shutil.rmtree(temp_dir, ignore_errors=True)
            return None

        # 打开文档
        doc = Document(temp_docx_path)

        # 创建图片目录
        images_dir = os.path.join(output_dir, "images")
        os.makedirs(images_dir, exist_ok=True)

        # 准备文档结构
        document_structure = {
            "metadata": extract_metadata(doc, docx_path),
            "sections": []
        }

        # 图片引用字典
        image_references = {}

        # 提取页眉页脚中的图片
        extract_header_footer_images(doc, images_dir, image_references)

        # 创建根节点
        root_section = {
            "type": "section",
            "title": "根节点",
            "level": 0,
            "content": []
        }
        document_structure["sections"].append(root_section)

        # 使用栈管理标题层级
        stack = deque([root_section])

        # 状态跟踪
        in_toc = False  # 是否在目录部分
        list_counter = defaultdict(int)  # 多级列表计数器
        table_counter = 0  # 表格计数器

        # 遍历文档块
        for block in iter_block_items(doc):
            # 段落处理
            if isinstance(block, Paragraph):
                text = clean_text(block.text)

                # 检测目录开始
                if not in_toc and re.match(r'^(目\s*录|contents?)$', text, re.IGNORECASE):
                    in_toc = True
                    continue

                # 检测目录结束
                if in_toc:
                    if get_heading_level(block) > 0 or len(text) > 50:
                        in_toc = False
                    else:
                        continue  # 跳过目录内容

                # 标题处理
                heading_level = get_heading_level(block)
                if heading_level > 0:
                    # 创建新章节
                    new_section = {
                        "type": "section",
                        "title": text,
                        "level": heading_level,
                        "content": []
                    }

                    # 确定父章节
                    while stack and stack[-1]["level"] >= heading_level:
                        stack.pop()

                    if stack:  # 添加到父章节
                        stack[-1]["content"].append(new_section)
                    else:  # 添加到根章节
                        root_section["content"].append(new_section)

                    stack.append(new_section)
                    continue

                # 列表项处理
                if is_list_item(block):
                    list_level, prefix = get_list_info(block, list_counter)

                    # 提取列表项中的图片
                    image_nodes = extract_paragraph_images(block, images_dir, image_references)

                    # 创建列表项节点
                    list_item = {
                        "type": "list_item",
                        "text": prefix + text,
                        "level": list_level
                    }

                    # 添加到当前章节
                    if stack:
                        stack[-1]["content"].append(list_item)

                    # 添加图片作为独立节点
                    if image_nodes:
                        for img in image_nodes:
                            stack[-1]["content"].append({
                                "type": "image",
                                "url": img["url"],
                                "format": img["format"],
                                "width": img["width"],
                                "height": img["height"],
                                "size": img["size"]
                            })
                    continue

                # 普通段落处理
                if text or block.runs:
                    # 提取段落中的图片
                    image_nodes = extract_paragraph_images(block, images_dir, image_references)

                    # 添加文本段落
                    if text:
                        para_item = {
                            "type": "paragraph",
                            "text": text,
                            "bold": any(run.bold for run in block.runs) if hasattr(block, 'runs') else False,
                            "italic": any(run.italic for run in block.runs) if hasattr(block, 'runs') else False
                        }

                        # 添加到当前章节
                        if stack:
                            stack[-1]["content"].append(para_item)

                    # 添加图片作为独立节点
                    if image_nodes:
                        for img in image_nodes:
                            stack[-1]["content"].append({
                                "type": "image",
                                "url": img["url"],
                                "format": img["format"],
                                "width": img["width"],
                                "height": img["height"],
                                "size": img["size"]
                            })

            # 表格处理
            elif isinstance(block, Table):
                table_data = parse_table(block, table_counter, image_references, images_dir)
                table_item = {
                    "type": "table",
                    "index": table_counter,
                    "rows": table_data
                }
                table_counter += 1

                # 添加到当前章节
                if stack:
                    stack[-1]["content"].append(table_item)

        # 添加页眉页脚图片
        header_footer_images = []
        for section in doc.sections:
            if hasattr(section, 'header_part') and hasattr(section.header_part, 'images'):
                header_footer_images.extend(section.header_part.images)
            if hasattr(section, 'footer_part') and hasattr(section.footer_part, 'images'):
                header_footer_images.extend(section.footer_part.images)

        if header_footer_images:
            document_structure["header_footer_images"] = [
                {
                    "type": "image",
                    "url": img["url"],
                    "format": img["format"],
                    "width": img["width"],
                    "height": img["height"],
                    "size": img["size"],
                    "location": "header" if "header" in img.get("context", "") else "footer"
                }
                for img in header_footer_images
            ]

        # 添加图片引用
        if image_references:
            document_structure["images"] = image_references
        else:
            logger.warning(f"文档 {os.path.basename(docx_path)} 中没有检测到图片")

        # 清理临时目录
        shutil.rmtree(temp_dir, ignore_errors=True)
        logger.info(f"清理临时目录: {temp_dir}")

        return document_structure

    except Exception as e:
        logger.error(f"解析文档 {docx_path} 失败: {e}")
        logger.error(traceback.format_exc())
        return None


def safe_filename(filename):
    """创建安全的文件名，移除无效字符"""
    # 移除文件系统不允许的字符
    safe_name = re.sub(r'[<>:"/\\|?*\[\]]', '_', filename)
    # 限制长度
    return safe_name[:150]


def process_docx_folder(input_folder, output_base_dir):
    """
    批量处理文件夹中的所有DOCX文件
    """
    # 确保输出目录存在
    os.makedirs(output_base_dir, exist_ok=True)

    # 准备汇总数据
    all_documents = []

    # 查找所有DOCX文件
    docx_files = [f for f in os.listdir(input_folder) if f.lower().endswith('.docx')]

    if not docx_files:
        logger.warning(f"在 {input_folder} 中没有找到DOCX文件")
        return

    logger.info(f"开始处理 {len(docx_files)} 个DOCX文件...")

    processed_count = 0
    failed_files = []

    for filename in docx_files:
        docx_path = os.path.join(input_folder, filename)
        logger.info(f"处理文件: {filename}")

        # 为每个文件创建输出目录
        safe_name = safe_filename(filename.replace('.docx', ''))
        output_dir = os.path.join(output_base_dir, safe_name)
        os.makedirs(output_dir, exist_ok=True)

        # 解析文档
        document_structure = parse_docx(docx_path, output_dir)

        if not document_structure:
            logger.error(f"跳过 {filename}，解析失败")
            failed_files.append(filename)
            continue

        processed_count += 1

        # 保存为JSON文件
        json_path = os.path.join(output_dir, "document.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(document_structure, f, ensure_ascii=False, indent=2)

        # 添加到汇总列表
        all_documents.append({
            "file": filename,
            "path": json_path
        })

        # 统计图片数量
        total_images = len(document_structure.get("images", {}))
        hf_images = len(document_structure.get("header_footer_images", []))
        logger.info(f"文件 {filename} 处理完成，检测到 {total_images} 张正文图片和 {hf_images} 张页眉页脚图片")

    # 保存汇总信息
    summary_path = os.path.join(output_base_dir, "summary.json")
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump({
            "input_folder": input_folder,
            "total_files": len(docx_files),
            "processed": processed_count,
            "failed": len(failed_files),
            "failed_files": failed_files,
            "documents": all_documents
        }, f, ensure_ascii=False, indent=2)

    logger.info(f"处理完成! 成功处理 {processed_count}/{len(docx_files)} 个文件")
    logger.info(f"汇总信息保存在 {summary_path}")
    return processed_count


if __name__ == "__main__":
    # 示例使用 - 批量处理文件夹
    input_folder = "Docx_parser/Files/PLM2.0"  # DOCX文件夹路径
    output_base_dir = "Docx_parser/Files/structured_docs"  # 输出目录

    # 检查输入目录是否存在
    if not os.path.exists(input_folder):
        logger.error(f"输入目录不存在: {input_folder}")
        sys.exit(1)

    # 处理所有DOCX文件
    processed_count = process_docx_folder(input_folder, output_base_dir)

    # 打印总结报告
    if processed_count > 0:
        summary_path = os.path.join(output_base_dir, "summary.json")
        if os.path.exists(summary_path):
            with open(summary_path, "r", encoding="utf-8") as f:
                summary = json.load(f)
                print("\n处理总结报告:")
                print(f"处理文件夹: {summary['input_folder']}")
                print(f"文件总数: {summary['total_files']}")
                print(f"成功处理: {summary['processed']}")
                print(f"失败文件: {summary['failed']}")
                if summary['failed_files']:
                    print("失败文件列表:")
                    for fname in summary['failed_files']:
                        print(f"  - {fname}")
